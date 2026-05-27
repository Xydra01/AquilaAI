import os
import base64
import csv
import io
import mimetypes
from pathlib import Path

from pdf_text import configure_mupdf_quiet, extract_pdf_text

configure_mupdf_quiet()

import docx

MAX_FILE_BYTES = 5 * 1024 * 1024
MAX_CHARS_PER_CHUNK = 90000
MAX_CSV_ROWS = 200

# Code / config / markup — parsed as UTF-8 text in all modes (chat, tasks, learn, etc.)
ATTACHMENT_CODE_EXTENSIONS = frozenset({
    ".py", ".pyw", ".pyi", ".ipynb",
    ".js", ".mjs", ".cjs", ".jsx", ".ts", ".tsx",
    ".java", ".kt", ".kts", ".scala", ".groovy",
    ".c", ".cc", ".cpp", ".cxx", ".h", ".hh", ".hpp", ".hxx",
    ".cs", ".vb", ".fs", ".fsx",
    ".go", ".rs", ".zig", ".nim", ".d",
    ".rb", ".rake", ".erb",
    ".php", ".phtml",
    ".swift", ".m", ".mm", ".mlx",
    ".r", ".R", ".jl",
    ".f", ".f90", ".f95", ".for", ".f77",
    ".lua", ".pl", ".pm", ".rkt", ".scm", ".ss", ".clj", ".cljs", ".cljc",
    ".hs", ".lhs", ".ml", ".mli", ".elm", ".ex", ".exs", ".erl", ".hrl",
    ".dart", ".v", ".sv", ".svh", ".vhd", ".vhdl",
    ".asm", ".s", ".S",
    ".sh", ".bash", ".zsh", ".fish", ".ps1", ".psm1", ".bat", ".cmd",
    ".sql", ".psql",
    ".html", ".htm", ".xhtml", ".css", ".scss", ".sass", ".less",
    ".vue", ".svelte", ".astro",
    ".xml", ".xsl", ".xslt", ".svg",
    ".json", ".jsonc", ".json5", ".jsonl", ".ndjson",
    ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf", ".config",
    ".env", ".properties",
    ".cmake", ".gradle", ".groovy", ".gradle.kts",
    ".tex", ".bib", ".sty",
    ".proto", ".thrift",
    ".gitignore", ".gitattributes", ".editorconfig",
    ".graphql", ".gql",
    ".make", ".mk",
    ".log",
})

# General text / docs (non-code-specific labels in file dialog)
ATTACHMENT_DOC_EXTENSIONS = frozenset({
    ".txt", ".md", ".markdown", ".rst", ".adoc", ".org",
    ".csv", ".tsv",
})

TEXT_EXTENSIONS = ATTACHMENT_CODE_EXTENSIONS | ATTACHMENT_DOC_EXTENSIONS

# Extensionless files still treated as text when attached by basename
EXTENSIONLESS_TEXT_NAMES = frozenset({
    "dockerfile",
    "makefile",
    "gnumakefile",
    "cmakelists.txt",
    "gemfile",
    "procfile",
    "vagrantfile",
    "brewfile",
    "rakefile",
    "justfile",
    "license",
    "readme",
})

# Register common code MIME guesses (Windows often leaves these unset)
for _ext, _mime in (
    (".m", "text/plain"),
    (".mlx", "application/octet-stream"),
    (".r", "text/plain"),
    (".jl", "text/plain"),
    (".ipynb", "application/json"),
    (".f90", "text/plain"),
    (".vhd", "text/plain"),
):
    mimetypes.add_type(_mime, _ext)


def is_attachment_text_path(path: str | Path) -> bool:
    """True if this path should be parsed as UTF-8 attachment text."""
    p = Path(path)
    ext = p.suffix.lower()
    if ext in TEXT_EXTENSIONS:
        return True
    return p.name.lower() in EXTENSIONLESS_TEXT_NAMES


def attachment_dialog_filter() -> str:
    """Qt QFileDialog filter string for workspace attachments."""
    code_globs = " ".join(f"*{e}" for e in sorted(ATTACHMENT_CODE_EXTENSIONS))
    doc_globs = " ".join(f"*{e}" for e in sorted(ATTACHMENT_DOC_EXTENSIONS))
    return (
        "All Files (*);;"
        f"Code & scripts ({code_globs});;"
        f"Documents ({doc_globs} *.pdf *.docx *.html *.htm);;"
        "MATLAB & scientific (*.m *.mlx *.r *.R *.jl *.f *.f90 *.f95);;"
        "Images (*.png *.jpg *.jpeg *.webp *.gif);;"
        "PDFs (*.pdf);;"
        "Spreadsheets (*.csv *.tsv *.xlsx)"
    )


def _read_bytes(path: str) -> tuple[bytes, str] | None:
    if not os.path.exists(path):
        return None
    file_name = os.path.basename(path)
    size = os.path.getsize(path)
    if size > MAX_FILE_BYTES:
        return None
    with open(path, "rb") as f:
        return f.read(), file_name


def _wrap_text(file_name: str, text: str) -> str:
    return f"--- START FILE: {file_name} ---\n{text}\n--- END FILE ---"


def _parse_image(file_bytes: bytes, mime_type: str) -> dict:
    b64 = base64.b64encode(file_bytes).decode("utf-8")
    return {"type": "image", "payload": b64, "mime": mime_type or "image/jpeg"}


def _parse_pdf(file_bytes: bytes, file_name: str) -> str:
    extracted = extract_pdf_text(file_bytes)
    if not extracted.strip():
        return _wrap_text(
            file_name,
            "[System: PDF opened but no extractable text was found "
            "(may be scan-only or heavily image-based).]",
        )
    return _wrap_text(file_name, extracted)


def _parse_docx(file_bytes: bytes, file_name: str) -> str:
    document = docx.Document(io.BytesIO(file_bytes))
    parts = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in document.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                if cell.text.strip() and cell.text.strip() not in row_data:
                    row_data.append(cell.text.strip())
            if row_data:
                parts.append(" | ".join(row_data))
    return _wrap_text(file_name, "\n".join(parts))


def _parse_csv(file_bytes: bytes, file_name: str) -> str:
    text = file_bytes.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    lines = []
    for i, row in enumerate(reader):
        if i >= MAX_CSV_ROWS:
            lines.append(f"... [truncated after {MAX_CSV_ROWS} rows]")
            break
        lines.append(", ".join(row))
    return _wrap_text(file_name, "\n".join(lines))


def _parse_xlsx(file_bytes: bytes, file_name: str) -> str:
    try:
        import openpyxl
    except ImportError:
        return (
            f"[System: Cannot parse {file_name} — install openpyxl for Excel support]"
        )
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        parts.append(f"## Sheet: {sheet.title}")
        for i, row in enumerate(sheet.iter_rows(values_only=True)):
            if i >= MAX_CSV_ROWS:
                parts.append(f"... [truncated after {MAX_CSV_ROWS} rows]")
                break
            parts.append(", ".join(str(c) if c is not None else "" for c in row))
    wb.close()
    return _wrap_text(file_name, "\n".join(parts))


def _parse_html(file_bytes: bytes, file_name: str) -> str:
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        return file_bytes.decode("utf-8", errors="replace")
    soup = BeautifulSoup(file_bytes.decode("utf-8", errors="replace"), "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    return _wrap_text(file_name, soup.get_text(separator="\n", strip=True))


def _parse_mlx(file_bytes: bytes, file_name: str) -> str:
    """MATLAB live script (.mlx) is a ZIP; extract readable XML/text when possible."""
    import zipfile

    try:
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            names = zf.namelist()
            parts = [f"[MATLAB live script: {file_name}]"]
            for name in names:
                if name.endswith(".xml") or name.endswith(".m"):
                    try:
                        raw = zf.read(name).decode("utf-8", errors="replace")
                        if raw.strip():
                            parts.append(f"## {name}\n{raw[:50000]}")
                    except Exception:
                        continue
            if len(parts) > 1:
                return _wrap_text(file_name, "\n\n".join(parts))
    except Exception as e:
        return _wrap_text(
            file_name,
            f"[System: Could not unpack MLX {file_name} — {e}. "
            "Export as .m or paste script text.]",
        )
    return _wrap_text(
        file_name,
        "[System: MLX opened but no extractable text; attach .m script instead.]",
    )


def _parse_text(file_bytes: bytes, file_name: str) -> str:
    return _wrap_text(file_name, file_bytes.decode("utf-8", errors="replace"))


def _dispatch_parse(path: str, file_bytes: bytes, file_name: str, mime_type: str | None):
    ext = Path(file_name).suffix.lower()
    base_lower = Path(file_name).name.lower()

    if mime_type and mime_type.startswith("image/"):
        return _parse_image(file_bytes, mime_type)

    if mime_type == "application/pdf" or ext == ".pdf":
        try:
            return {"type": "text", "payload": _parse_pdf(file_bytes, file_name)}
        except Exception as e:
            return {"type": "text", "payload": f"[System: Could not parse PDF {file_name} - {e}]"}

    if ext == ".docx":
        try:
            return {"type": "text", "payload": _parse_docx(file_bytes, file_name)}
        except Exception as e:
            return {"type": "text", "payload": f"[System: Could not parse DOCX {file_name} - {e}]"}

    if ext in (".csv", ".tsv"):
        try:
            return {"type": "text", "payload": _parse_csv(file_bytes, file_name)}
        except Exception as e:
            return {"type": "text", "payload": f"[System: Could not parse CSV {file_name} - {e}]"}

    if ext == ".xlsx":
        return {"type": "text", "payload": _parse_xlsx(file_bytes, file_name)}

    if ext == ".mlx":
        return {"type": "text", "payload": _parse_mlx(file_bytes, file_name)}

    if ext in (".html", ".htm") or (mime_type and "html" in (mime_type or "")):
        return {"type": "text", "payload": _parse_html(file_bytes, file_name)}

    if is_attachment_text_path(path) or is_attachment_text_path(file_name):
        return {"type": "text", "payload": _parse_text(file_bytes, file_name)}

    if mime_type and mime_type.startswith("text/"):
        return {"type": "text", "payload": _parse_text(file_bytes, file_name)}

    try:
        return {"type": "text", "payload": _parse_text(file_bytes, file_name)}
    except Exception:
        return {
            "type": "text",
            "payload": f"[System: Unsupported or binary file type for {file_name}]",
        }


def process_local_attachments(file_paths):
    """Parse local files into text chunks and base64 image payloads for the agent."""
    text_payloads = []
    image_payloads = []

    for path in file_paths:
        result = _read_bytes(path)
        if result is None:
            file_name = os.path.basename(path)
            if os.path.exists(path) and os.path.getsize(path) > MAX_FILE_BYTES:
                text_payloads.append(
                    f"[System: File {file_name} exceeds {MAX_FILE_BYTES // (1024*1024)}MB limit and was skipped]"
                )
            continue

        file_bytes, file_name = result
        mime_type, _ = mimetypes.guess_type(path)
        parsed = _dispatch_parse(path, file_bytes, file_name, mime_type)

        if parsed["type"] == "image":
            image_payloads.append(parsed["payload"])
        else:
            text_payloads.append(parsed["payload"])

    combined_text = "\n\n".join(text_payloads)
    text_chunks = []
    if combined_text.strip():
        for i in range(0, len(combined_text), MAX_CHARS_PER_CHUNK):
            text_chunks.append(combined_text[i : i + MAX_CHARS_PER_CHUNK])

    return text_chunks, image_payloads


def _strip_file_wrapper(text: str) -> str:
    """Remove --- START/END FILE --- wrappers from parsed attachment text."""
    raw = (text or "").strip()
    if "--- START FILE:" not in raw:
        return raw
    if "--- END FILE ---" in raw:
        inner = raw.split("--- START FILE:", 1)[-1]
        inner = inner.split("\n", 1)[-1] if "\n" in inner else inner
        inner = inner.rsplit("--- END FILE ---", 1)[0]
        return inner.strip()
    return raw


def extract_indexable_text(path: str | Path) -> str:
    """Extract plain text from a local file for Learn archive/course indexing."""
    path = Path(path)
    result = _read_bytes(str(path))
    if result is None:
        return ""
    file_bytes, file_name = result
    mime_type, _ = mimetypes.guess_type(str(path))
    parsed = _dispatch_parse(str(path), file_bytes, file_name, mime_type)
    if parsed.get("type") != "text":
        return ""
    return _strip_file_wrapper(parsed.get("payload", ""))
